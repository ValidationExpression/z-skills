#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from apply_format import load_config  # noqa: E402
from gen_from_template import build_document  # noqa: E402


CONTENT = {
    "cover": {
        "title": ["测试公司", "示例项目"],
        "document_type": "报告",
        "metadata": ["编制单位：测试公司", "日期：2026年8月"],
    },
    "blocks": [
        {"type": "heading", "level": 1, "text": "第一章 项目概述"},
        {"type": "paragraph", "text": '他说"开始执行"。'},
        {
            "type": "table",
            "headers": ["字段", "内容"],
            "rows": [["项目", "示例项目"], ["状态", "进行中"]],
        },
    ],
}


class GenFromTemplateTests(unittest.TestCase):
    def test_load_config_does_not_leak_between_runs(self):
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "custom.json"
            config_path.write_text(
                json.dumps({"headings": {"Heading 1": {"size_pt": 28}}}),
                encoding="utf-8",
            )
            self.assertEqual(load_config(config_path)["headings"]["Heading 1"]["size_pt"], 28)
            self.assertEqual(load_config(None)["headings"]["Heading 1"]["size_pt"], 15)

    def test_build_document_uses_fingerprint_factories(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "generated.docx"
            build_document(CONTENT, load_config(None), output)
            self.assertTrue(output.exists())

            doc = Document(output)
            heading = next(p for p in doc.paragraphs if p.text == "第一章 项目概述")
            page_break = heading._p.pPr.find(qn("w:pageBreakBefore"))
            self.assertIsNotNone(page_break)

            body = next(p for p in doc.paragraphs if "开始执行" in p.text)
            self.assertEqual(body.text, "他说“开始执行”。")
            self.assertEqual(body.paragraph_format.first_line_indent.emu, 304800)
            self.assertEqual(body.paragraph_format.line_spacing, 1.5)
            run = body.runs[0]
            self.assertEqual(run.font.name, "仿宋")
            self.assertEqual(
                run.font.element.rPr.rFonts.get(qn("w:eastAsia")), "仿宋"
            )
            self.assertEqual(run.font.size.pt, 12)

            borders = doc.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
            self.assertIsNotNone(borders)
            self.assertTrue(doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0].bold)

            verification = subprocess.run(
                [sys.executable, str(SCRIPTS / "verify_format.py"), str(output)],
                text=True,
                capture_output=True,
                check=False,
            )
            self.assertEqual(verification.returncode, 0, verification.stdout)

    def test_cli_accepts_extracted_style_config(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            content_path = tmp_path / "content.json"
            config_path = tmp_path / "fingerprint.json"
            output_path = tmp_path / "cli.docx"
            content_path.write_text(
                json.dumps(CONTENT, ensure_ascii=False), encoding="utf-8"
            )
            config_path.write_text(
                json.dumps(
                    {
                        "body": {
                            "font": ["宋体", "宋体"],
                            "size_pt": 11,
                            "first_line_indent_emu": 279400,
                            "line_spacing": "1.25",
                        },
                        "headings": {"Heading 1": {"size_pt": 16}},
                        "tables": {"cell_size_pt": 10},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPTS / "gen_from_template.py"),
                    str(content_path),
                    "--config",
                    str(config_path),
                    "--out",
                    str(output_path),
                ],
                text=True,
                capture_output=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertTrue(output_path.exists())
            doc = Document(output_path)
            body = next(p for p in doc.paragraphs if "开始执行" in p.text)
            self.assertEqual(body.runs[0].font.name, "宋体")
            self.assertEqual(body.runs[0].font.size.pt, 11)


if __name__ == "__main__":
    unittest.main()
